"""DOCX report generator matching PRASHO's report format.

Generates a Word document with:
- Title page
- Executive summary
- Purpose and methodology sections
- Per-component analysis with enabler/barrier tables, bar charts, radar chart, heatmap
- AI-generated recommendations
- Appendix with all classified remarks
"""

import io
import logging
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from core.charts import (
    create_radar_chart, create_theme_heatmap,
    create_enabler_barrier_bar, create_component_bar,
    build_theme_data_from_analyses,
    COMPONENT_ORDER,
)

logger = logging.getLogger("sehra.report_gen")


COMPONENT_DISPLAY_NAMES = {
    "context": "Context",
    "policy": "Sectoral Legislation, Policy and Strategy",
    "service_delivery": "Institutional and Service Delivery Environment",
    "human_resources": "Human Resources",
    "supply_chain": "Supply Chain",
    "barriers": "Barriers",
}

# Colors matching SEHRA branding
COLORS = {
    "enabler": RGBColor(0x0D, 0x73, 0x77),  # Teal
    "barrier": RGBColor(0xCC, 0x33, 0x33),   # Red
    "header_bg": RGBColor(0x0D, 0x73, 0x77),
}


def generate_bar_chart(enabler_count: int, barrier_count: int,
                       component_name: str) -> io.BytesIO:
    """Generate a bar chart comparing enablers vs barriers."""
    fig, ax = plt.subplots(figsize=(6, 3))

    categories = ['Enablers', 'Barriers']
    values = [enabler_count, barrier_count]
    colors = ['#0D7377', '#CC3333']

    bars = ax.bar(categories, values, color=colors, width=0.5)

    # Add value labels on bars
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                str(val), ha='center', va='bottom', fontweight='bold', fontsize=12)

    ax.set_title(f'{component_name}', fontsize=13, fontweight='bold', pad=10)
    ax.set_ylabel('Count', fontsize=11)
    ax.set_ylim(0, max(values) * 1.3 if max(values) > 0 else 5)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_overall_chart(components_data: list[dict]) -> io.BytesIO:
    """Generate an overall bar chart showing all components."""
    fig, ax = plt.subplots(figsize=(10, 5))

    names = [d["name"] for d in components_data]
    enablers = [d["enabler_count"] for d in components_data]
    barriers = [d["barrier_count"] for d in components_data]

    x = np.arange(len(names))
    width = 0.35

    bars1 = ax.bar(x - width/2, enablers, width, label='Enablers', color='#0D7377')
    bars2 = ax.bar(x + width/2, barriers, width, label='Barriers', color='#CC3333')

    for bar, val in zip(bars1, enablers):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                    str(val), ha='center', va='bottom', fontsize=9, fontweight='bold')
    for bar, val in zip(bars2, barriers):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
                    str(val), ha='center', va='bottom', fontsize=9, fontweight='bold')

    ax.set_ylabel('Count', fontsize=11)
    ax.set_title('Enablers vs Barriers by Component', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=20, ha='right', fontsize=9)
    ax.legend()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


# XML 1.0 illegal characters (control chars except \t \n \r)
_ILLEGAL_CHARS_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def _clean(value: str) -> str:
    """Strip characters that are illegal in XML (used by DOCX)."""
    return _ILLEGAL_CHARS_RE.sub('', value) if isinstance(value, str) else str(value)


def _add_styled_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '0D7377',
        })
        shading.append(shading_elem)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = _clean(str(cell_text))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def generate_report(sehra_data: dict, component_analyses: list[dict],
                    header_info: dict,
                    executive_summary: str = "",
                    recommendations: str = "",
                    generated_at_ist: str = "",
                    requester_ip: str = "",
                    exported_by: str = "") -> io.BytesIO:
    """Generate a complete DOCX report.

    Args:
        sehra_data: SEHRA record data
        component_analyses: List of component analysis dicts from DB
        header_info: {country, district, assessment_date}
        executive_summary: Optional AI-generated executive summary
        recommendations: Optional AI-generated recommendations
        generated_at_ist: IST-formatted generation timestamp
        requester_ip: IP address of the requester
        exported_by: Username of the exporter

    Returns:
        BytesIO containing the DOCX file
    """
    logger.info("Generating DOCX report for %s", header_info.get("country", ""))
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Analysis Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLORS["header_bg"]

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    country = header_info.get("country", "")
    run = subtitle.add_run(f"{country} Scoping Module")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COLORS["header_bg"]

    district = header_info.get("district", "")
    if district:
        dist_p = doc.add_paragraph()
        dist_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = dist_p.add_run(f"District: {district}")
        run.font.size = Pt(14)

    date_str = header_info.get("assessment_date", "")
    if date_str:
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(f"Date: {date_str}")
        run.font.size = Pt(12)

    doc.add_paragraph("")
    brand_p = doc.add_paragraph()
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand_p.add_run("SEHRA Analyzer — PRASHO Foundation")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # --- EXECUTIVE SUMMARY ---
    if executive_summary:
        doc.add_heading("Executive Summary", level=1)
        for para in executive_summary.strip().split("\n\n"):
            if para.strip():
                doc.add_paragraph(_clean(para.strip()))
        doc.add_page_break()

    # --- PURPOSE ---
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "The SEHRA Scoping Module provides a rapid overview of the policy, strategy, "
        "institutional and service delivery environment for a school eye health programme. "
        "This report presents the analysis of the scoping module data collected in "
        f"{country}{' (' + district + ')' if district else ''}."
    )

    # --- METHODOLOGY ---
    doc.add_heading("Methodology", level=1)

    doc.add_heading("Numeric data", level=2)
    doc.add_paragraph(
        "The scoping module questionnaire contains both numeric and textual data. "
        "Numeric responses (Yes/No) were coded as binary variables: Yes = 1 (Enabler) "
        "and No = 0 (Barrier). Some items are reverse-scored where Yes indicates a barrier. "
        "Enabler and barrier counts were calculated per component."
    )

    doc.add_heading("Textual data - Development of Thematic Framework", level=2)
    doc.add_paragraph(
        "Free-text remarks were analyzed using a deductive thematic framework with "
        "11 cross-cutting themes. Each remark was classified as an enabler, barrier, "
        "strength, or weakness and mapped to one or more themes. AI-assisted analysis "
        "(Claude, Anthropic) was used with human review to ensure accuracy."
    )

    # List themes
    doc.add_paragraph("The 11 cross-cutting themes used for analysis:")
    themes = [
        "Institutional Structure and Stakeholders",
        "Operationalization Strategies",
        "Coordination and Integration",
        "Funding",
        "Local Capacity and Service Delivery",
        "Accessibility and Inclusivity",
        "Cost, Availability and Affordability",
        "Data Considerations",
        "Sociocultural Factors and Compliance",
        "Services at Higher Levels of Health System",
        "Provision of Eyeglasses",
    ]
    for i, theme in enumerate(themes, 1):
        doc.add_paragraph(f"{i}. {theme}", style='List Number')

    doc.add_page_break()

    # --- RESULTS ---
    doc.add_heading("Results", level=1)

    # Build component data
    comp_lookup = {ca["component"]: ca for ca in component_analyses}
    comp_scores = {}
    components_chart_data = []

    for comp_key in COMPONENT_ORDER:
        ca = comp_lookup.get(comp_key)
        if ca:
            components_chart_data.append({
                "name": COMPONENT_DISPLAY_NAMES.get(comp_key, comp_key),
                "enabler_count": ca.get("enabler_count", 0),
                "barrier_count": ca.get("barrier_count", 0),
            })
            comp_scores[comp_key] = {
                "enabler_count": ca.get("enabler_count", 0),
                "barrier_count": ca.get("barrier_count", 0),
            }

    # Radar chart
    try:
        _, radar_png = create_radar_chart(comp_scores)
        if radar_png:
            doc.add_paragraph("Readiness Profile:")
            doc.add_picture(io.BytesIO(radar_png), width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
    except Exception as e:
        logger.warning("Failed to add radar chart to DOCX: %s", e)

    # Overall chart (matplotlib version for DOCX)
    if components_chart_data:
        overall_chart = generate_overall_chart(components_chart_data)
        doc.add_picture(overall_chart, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    # Theme heatmap
    try:
        theme_data = build_theme_data_from_analyses(component_analyses)
        if theme_data:
            _, heatmap_png = create_theme_heatmap(theme_data)
            if heatmap_png:
                doc.add_paragraph("Theme Distribution:")
                doc.add_picture(io.BytesIO(heatmap_png), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
    except Exception as e:
        logger.warning("Failed to add heatmap to DOCX: %s", e)

    # --- PER-COMPONENT SECTIONS ---
    for comp_key in COMPONENT_ORDER:
        ca = comp_lookup.get(comp_key)
        if not ca:
            continue

        comp_name = COMPONENT_DISPLAY_NAMES.get(comp_key, comp_key)
        doc.add_heading(comp_name, level=2)

        enabler_count = ca.get("enabler_count", 0)
        barrier_count = ca.get("barrier_count", 0)

        doc.add_paragraph(
            f"Quantitative analysis: {enabler_count} Enablers, {barrier_count} Barriers"
        )

        # Bar chart
        chart_buf = generate_bar_chart(enabler_count, barrier_count, comp_name)
        doc.add_picture(chart_buf, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Scored items (codebook answers — always populated)
        items = ca.get("items", [])
        scored_enablers = [i for i in items if i.get("classification") == "enabler"]
        scored_barriers = [i for i in items if i.get("classification") == "barrier"]

        report_sections = ca.get("report_sections", {})

        if scored_enablers:
            doc.add_heading(f"Enabler Items ({len(scored_enablers)})", level=3)
            enabler_summary = report_sections.get("enabler_summary", {})
            if enabler_summary and enabler_summary.get("content"):
                doc.add_paragraph(_clean(enabler_summary["content"]))
            rows = [[i.get("item_id", ""), i.get("question", "")[:200], str(i.get("answer", ""))] for i in scored_enablers]
            _add_styled_table(doc, ["Item", "Question", "Answer"], rows)

        if scored_barriers:
            doc.add_heading(f"Barrier Items ({len(scored_barriers)})", level=3)
            barrier_summary = report_sections.get("barrier_summary", {})
            if barrier_summary and barrier_summary.get("content"):
                doc.add_paragraph(_clean(barrier_summary["content"]))
            rows = [[i.get("item_id", ""), i.get("question", "")[:200], str(i.get("answer", ""))] for i in scored_barriers]
            _add_styled_table(doc, ["Item", "Question", "Answer"], rows)

        # AI-classified remarks (only when text remarks exist)
        qual_entries = ca.get("qualitative_entries", [])
        enabler_entries = [e for e in qual_entries if e.get("classification") in ("enabler", "strength")]
        barrier_entries = [e for e in qual_entries if e.get("classification") in ("barrier", "weakness")]

        if enabler_entries:
            doc.add_heading("Classified Enabler Remarks", level=3)
            theme_groups = {}
            for e in enabler_entries:
                theme = e.get("theme", "Other")
                if theme not in theme_groups:
                    theme_groups[theme] = []
                theme_groups[theme].append(e.get("remark_text", ""))
            rows = []
            for theme, remarks in theme_groups.items():
                summary = "; ".join(r for r in remarks if r)[:300]
                rows.append([theme, summary, ""])
            if rows:
                _add_styled_table(doc, ["Theme", "Remarks", "Action Points"], rows)

        if barrier_entries:
            doc.add_heading("Classified Barrier Remarks", level=3)
            theme_groups = {}
            for e in barrier_entries:
                theme = e.get("theme", "Other")
                if theme not in theme_groups:
                    theme_groups[theme] = []
                theme_groups[theme].append(e.get("remark_text", ""))
            rows = []
            for theme, remarks in theme_groups.items():
                summary = "; ".join(r for r in remarks if r)[:300]
                rows.append([theme, summary, ""])
            if rows:
                _add_styled_table(doc, ["Theme", "Remarks", "Action Points"], rows)

        # Action points from report sections
        action_section = report_sections.get("action_points", {})
        if action_section and action_section.get("content"):
            doc.add_heading("Action Points", level=3)
            doc.add_paragraph(_clean(action_section["content"]))

        doc.add_paragraph("")  # Spacing

    # --- RECOMMENDATIONS ---
    doc.add_page_break()
    doc.add_heading("Recommendations", level=1)

    if recommendations:
        # Use AI-generated recommendations
        for para in recommendations.strip().split("\n"):
            line = para.strip()
            if line:
                doc.add_paragraph(_clean(line))
    else:
        # Fallback to generic recommendations
        doc.add_paragraph(
            "Based on the analysis of the SEHRA scoping module data, the following "
            "recommendations are made for planning and implementing a school eye health "
            f"programme in {country}:"
        )
        doc.add_paragraph(
            "1. Address identified barriers in policy alignment and implementation",
            style='List Number'
        )
        doc.add_paragraph(
            "2. Strengthen coordination between health and education ministries",
            style='List Number'
        )
        doc.add_paragraph(
            "3. Invest in capacity building for human resources",
            style='List Number'
        )
        doc.add_paragraph(
            "4. Improve supply chain for eyeglasses and optical equipment",
            style='List Number'
        )
        doc.add_paragraph(
            "5. Address sociocultural barriers through community engagement",
            style='List Number'
        )

    # --- APPENDIX ---
    all_entries = []
    for comp_key in COMPONENT_ORDER:
        ca = comp_lookup.get(comp_key)
        if not ca:
            continue
        for e in ca.get("qualitative_entries", []):
            all_entries.append({
                "component": COMPONENT_DISPLAY_NAMES.get(comp_key, comp_key),
                **e,
            })

    if all_entries:
        doc.add_page_break()
        doc.add_heading("Appendix: All Classified Remarks", level=1)
        rows = []
        for e in all_entries:
            rows.append([
                e.get("component", ""),
                e.get("theme", ""),
                e.get("classification", ""),
                f"{e.get('confidence', 0):.0%}",
                (e.get("remark_text", "") or "")[:150],
            ])
        _add_styled_table(
            doc,
            ["Component", "Theme", "Classification", "Confidence", "Remark"],
            rows,
        )

    # --- FOOTER ---
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("SEHRA Analyzer — Built for PRASHO Foundation by Samanvay Foundation")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["header_bg"]
    run.bold = True

    meta_parts = []
    if generated_at_ist:
        meta_parts.append(f"Generated on {generated_at_ist} IST")
    if exported_by:
        meta_parts.append(f"Exported by {exported_by}")
    if requester_ip:
        meta_parts.append(f"IP: {requester_ip}")
    if meta_parts:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_p.add_run(" | ".join(meta_parts))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("DOCX report generated")
    return buf
