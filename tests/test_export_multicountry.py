"""Test report generation for multi-country scenarios."""
import io
import pytest
from unittest.mock import patch

from core.report_gen import generate_report, generate_bar_chart, generate_overall_chart
from core.report_html import generate_html_report
from core.charts import (
    create_radar_chart, create_theme_heatmap,
    create_enabler_barrier_bar, create_component_bar,
    build_theme_data_from_analyses,
)


def _make_component_analyses(country="Liberia", num_components=6):
    """Helper to build component analyses for a given country."""
    components = ["context", "policy", "service_delivery",
                  "human_resources", "supply_chain", "barriers"]
    analyses = []
    for i, comp in enumerate(components[:num_components]):
        analyses.append({
            "id": f"ca-{i}",
            "component": comp,
            "enabler_count": 5 + i,
            "barrier_count": 2 + i,
            "items": [
                {
                    "item_id": f"{comp[0].upper()}{j}",
                    "question": f"Question {j} for {comp}",
                    "answer": "yes" if j % 2 == 0 else "no",
                    "remark": f"Remark about {comp} item {j}." if j < 3 else "",
                    "score": 1 if j % 2 == 0 else 0,
                    "classification": "enabler" if j % 2 == 0 else "barrier",
                    "is_reverse": False,
                    "component": comp,
                }
                for j in range(5)
            ],
            "qualitative_entries": [
                {
                    "id": f"qe-{comp}-{k}",
                    "remark_text": f"Qualitative remark {k} for {comp}.",
                    "item_id": f"{comp[0].upper()}{k}",
                    "theme": "Institutional Structure and Stakeholders",
                    "classification": "enabler" if k == 0 else "barrier",
                    "confidence": 0.85,
                    "edited_by_human": False,
                }
                for k in range(2)
            ],
            "report_sections": {
                "enabler_summary": {
                    "id": f"rs-e-{comp}",
                    "content": f"Enablers found in {comp}: strong institutional support.",
                    "edited_by_human": False,
                },
                "barrier_summary": {
                    "id": f"rs-b-{comp}",
                    "content": f"Barriers found in {comp}: limited resources.",
                    "edited_by_human": False,
                },
            },
        })
    return analyses


class TestReportMultiCountry:
    """Test that reports are correctly generated for different countries."""

    def test_docx_report_includes_country_name(self):
        """DOCX report title includes country name."""
        analyses = _make_component_analyses("Kenya")
        header = {"country": "Kenya", "district": "Nairobi",
                  "assessment_date": "2024-06-01"}
        sehra_data = {"id": "test-kenya", "country": "Kenya"}

        buf = generate_report(sehra_data, analyses, header)
        assert isinstance(buf, io.BytesIO)

        from docx import Document
        buf.seek(0)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Kenya" in full_text

    def test_docx_report_for_liberia(self, sample_component_analyses):
        """DOCX report generates correctly for Liberia data."""
        header = {"country": "Liberia", "district": "Montserrado",
                  "assessment_date": "2024-05-17"}
        sehra_data = {"id": "test-liberia", "country": "Liberia"}

        buf = generate_report(sehra_data, sample_component_analyses, header)
        buf.seek(0)
        assert buf.read(2) == b'PK'  # Valid DOCX zip signature

    def test_xlsx_report_country_context(self):
        """Report generation with all components for a non-Liberia country."""
        analyses = _make_component_analyses("Uganda")
        header = {"country": "Uganda", "district": "Kampala",
                  "assessment_date": "2024-07-15"}
        sehra_data = {"id": "test-uganda", "country": "Uganda"}

        buf = generate_report(sehra_data, analyses, header)
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(2) == b'PK'

    def test_html_report_country_context(self):
        """HTML report includes country in metadata."""
        analyses = _make_component_analyses("India")
        header = {"country": "India", "district": "Mumbai",
                  "assessment_date": "2024-08-01"}

        html = generate_html_report(analyses, header)
        assert isinstance(html, str)
        assert "India" in html
        assert "SEHRA" in html
        assert "<html" in html

    def test_report_with_empty_components(self):
        """Report generation handles components with zero items."""
        analyses = [
            {
                "id": "ca-empty",
                "component": "context",
                "enabler_count": 0,
                "barrier_count": 0,
                "items": [],
                "qualitative_entries": [],
                "report_sections": {},
            }
        ]
        header = {"country": "TestCountry", "district": "", "assessment_date": ""}
        sehra_data = {"id": "test-empty", "country": "TestCountry"}

        buf = generate_report(sehra_data, analyses, header)
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(2) == b'PK'

    def test_report_charts_generated(self):
        """Charts (radar, bar) are generated without errors."""
        # Radar chart
        scores = {
            "context": {"enabler_count": 8, "barrier_count": 4},
            "policy": {"enabler_count": 17, "barrier_count": 3},
            "service_delivery": {"enabler_count": 13, "barrier_count": 2},
            "human_resources": {"enabler_count": 2, "barrier_count": 3},
            "supply_chain": {"enabler_count": 27, "barrier_count": 5},
            "barriers": {"enabler_count": 10, "barrier_count": 38},
        }
        fig, png = create_radar_chart(scores)
        assert fig is not None

        # Enabler/barrier bar chart
        data = [
            {"name": comp.title(), "enabler_count": v["enabler_count"],
             "barrier_count": v["barrier_count"]}
            for comp, v in scores.items()
        ]
        fig2, png2 = create_enabler_barrier_bar(data)
        assert fig2 is not None

        # Component bar chart
        fig3, png3 = create_component_bar(10, 5, "Context")
        assert fig3 is not None

    def test_theme_heatmap_from_analyses(self, sample_component_analyses):
        """Theme heatmap builds from component analyses without error."""
        theme_data = build_theme_data_from_analyses(sample_component_analyses)
        assert isinstance(theme_data, dict)

        if theme_data:
            fig, png = create_theme_heatmap(theme_data)
            assert fig is not None

    def test_html_report_with_executive_summary(self):
        """HTML report includes executive summary when provided."""
        analyses = _make_component_analyses("Liberia")
        header = {"country": "Liberia", "district": "Montserrado",
                  "assessment_date": "2024-05-17"}

        html = generate_html_report(
            analyses, header,
            executive_summary="This is the executive summary for multi-country test.",
            recommendations="1. First recommendation\n2. Second recommendation",
        )
        assert "executive summary for multi-country test" in html
        assert "First recommendation" in html

    def test_docx_report_with_recommendations(self):
        """DOCX report includes recommendations when provided."""
        analyses = _make_component_analyses("Kenya")
        header = {"country": "Kenya", "district": "Nairobi",
                  "assessment_date": "2024-06-01"}
        sehra_data = {"id": "test-recs", "country": "Kenya"}

        buf = generate_report(
            sehra_data, analyses, header,
            executive_summary="Executive summary content.",
            recommendations="1. Address funding gaps\n2. Strengthen coordination",
        )

        from docx import Document
        buf.seek(0)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in full_text

    def test_overall_bar_chart_returns_png(self):
        """Overall bar chart generates valid PNG."""
        data = [
            {"name": "Context", "enabler_count": 8, "barrier_count": 4},
            {"name": "Policy", "enabler_count": 17, "barrier_count": 3},
            {"name": "Service Delivery", "enabler_count": 13, "barrier_count": 2},
        ]
        buf = generate_overall_chart(data)
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(4) == b'\x89PNG'

    def test_component_bar_chart_returns_png(self):
        """Individual component bar chart generates valid PNG."""
        buf = generate_bar_chart(15, 8, "Policy")
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(4) == b'\x89PNG'

    def test_html_report_contains_plotly(self):
        """HTML report includes Plotly interactive charts."""
        analyses = _make_component_analyses("Liberia")
        header = {"country": "Liberia", "district": "", "assessment_date": ""}

        html = generate_html_report(analyses, header)
        assert "plotly" in html.lower()

    def test_report_with_single_component(self):
        """Report generation works with only one component."""
        analyses = _make_component_analyses("Liberia", num_components=1)
        header = {"country": "Liberia", "district": "", "assessment_date": ""}
        sehra_data = {"id": "test-single", "country": "Liberia"}

        buf = generate_report(sehra_data, analyses, header)
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(2) == b'PK'
