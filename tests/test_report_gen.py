"""Tests for report generation modules."""

import io
import pytest
from unittest.mock import patch

from core.report_gen import generate_report, generate_bar_chart, generate_overall_chart
from core.report_html import generate_html_report
from core.charts import (
    create_radar_chart, create_theme_heatmap,
    create_enabler_barrier_bar, create_component_bar,
    build_theme_data_from_analyses,
)


class TestCharts:
    def test_radar_chart(self):
        scores = {
            "context": {"enabler_count": 8, "barrier_count": 4},
            "policy": {"enabler_count": 17, "barrier_count": 3},
        }
        fig, png = create_radar_chart(scores)
        assert fig is not None
        # PNG may be empty if kaleido not installed - that's ok

    def test_theme_heatmap(self):
        theme_data = {
            "Funding": {"context": 2, "policy": 5},
            "Data Considerations": {"context": 1},
        }
        fig, png = create_theme_heatmap(theme_data)
        assert fig is not None

    def test_enabler_barrier_bar(self):
        data = [
            {"name": "Context", "enabler_count": 8, "barrier_count": 4},
            {"name": "Policy", "enabler_count": 17, "barrier_count": 3},
        ]
        fig, png = create_enabler_barrier_bar(data)
        assert fig is not None

    def test_component_bar(self):
        fig, png = create_component_bar(10, 5, "Context")
        assert fig is not None

    def test_build_theme_data(self, sample_component_analyses):
        theme_data = build_theme_data_from_analyses(sample_component_analyses)
        assert isinstance(theme_data, dict)
        assert len(theme_data) > 0


class TestDOCXReport:
    def test_generates_valid_docx(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "Montserrado", "assessment_date": "2024-05-17"}
        sehra_data = {"id": "test-id", "country": "Liberia"}

        buf = generate_report(sehra_data, sample_component_analyses, header)

        assert isinstance(buf, io.BytesIO)
        # Check it's a valid DOCX (starts with PK zip signature)
        buf.seek(0)
        assert buf.read(2) == b'PK'

    def test_contains_country(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "Montserrado", "assessment_date": ""}
        sehra_data = {"id": "test-id", "country": "Liberia"}

        buf = generate_report(sehra_data, sample_component_analyses, header)

        # Read DOCX content
        from docx import Document
        buf.seek(0)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Liberia" in full_text

    def test_with_executive_summary(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "", "assessment_date": ""}
        sehra_data = {"id": "test-id", "country": "Liberia"}

        buf = generate_report(
            sehra_data, sample_component_analyses, header,
            executive_summary="This is the executive summary.",
            recommendations="1. First recommendation\n2. Second recommendation",
        )

        from docx import Document
        buf.seek(0)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in full_text
        assert "executive summary" in full_text.lower()


class TestBarChart:
    def test_returns_bytesio(self):
        buf = generate_bar_chart(10, 5, "Context")
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        # Check it's a valid PNG
        assert buf.read(4) == b'\x89PNG'

    def test_overall_chart(self):
        data = [
            {"name": "Context", "enabler_count": 8, "barrier_count": 4},
            {"name": "Policy", "enabler_count": 17, "barrier_count": 3},
        ]
        buf = generate_overall_chart(data)
        assert isinstance(buf, io.BytesIO)
        buf.seek(0)
        assert buf.read(4) == b'\x89PNG'


class TestHTMLReport:
    def test_generates_html(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "Montserrado", "assessment_date": "2024-05-17"}

        html = generate_html_report(sample_component_analyses, header)

        assert isinstance(html, str)
        assert "Liberia" in html
        assert "SEHRA" in html
        assert "<html" in html

    def test_contains_plotly(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "", "assessment_date": ""}

        html = generate_html_report(sample_component_analyses, header)

        assert "plotly" in html.lower()

    def test_with_summary(self, sample_component_analyses):
        header = {"country": "Liberia", "district": "", "assessment_date": ""}

        html = generate_html_report(
            sample_component_analyses, header,
            executive_summary="Test executive summary here.",
            recommendations="1. Test recommendation.",
        )

        assert "Test executive summary" in html
        assert "Test recommendation" in html
